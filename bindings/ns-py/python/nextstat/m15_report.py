"""Publishable ICH M15 report renderer.

This module consumes already-built M15 JSON artifacts and renders a combined
publishable report as Markdown, PDF, and DOCX. It does not execute models or
recompute any regulated evidence.
"""

from __future__ import annotations

import argparse
import datetime as _dt
import json
import textwrap
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Mapping, Sequence


_DOCX_ZIP_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
_FIXED_DT = _dt.datetime(1980, 1, 1, 0, 0, 0)

_COL = {
    "ink": "#111827",
    "muted": "#4B5563",
    "border": "#E5E7EB",
    "panel": "#F3F4F6",
}


def _require_matplotlib() -> None:
    try:
        import matplotlib  # noqa: F401
    except Exception as exc:  # pragma: no cover
        raise ImportError(
            "Missing dependency: matplotlib. Install via `pip install nextstat[viz]`."
        ) from exc


def _require_python_docx() -> None:
    try:
        import docx  # noqa: F401
    except Exception as exc:  # pragma: no cover
        raise ImportError(
            "Missing dependency: python-docx. Install via `pip install python-docx`."
        ) from exc


def _apply_pub_style() -> None:
    from . import report as _report_style
    import matplotlib as mpl

    _report_style._apply_pub_style()
    mpl.rcParams.update(
        {
            "figure.constrained_layout.use": True,
            "savefig.bbox": "tight",
            "axes.grid": False,
            "axes.spines.top": False,
            "axes.spines.right": False,
            "axes.spines.left": False,
            "axes.spines.bottom": False,
            "font.size": 10.0,
            "svg.fonttype": "none",
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
        }
    )


def _coerce_json_artifact(
    artifact_or_path: Mapping[str, Any] | str | Path,
    *,
    expected_schema_version: str,
) -> dict[str, Any]:
    if isinstance(artifact_or_path, Mapping):
        doc = dict(artifact_or_path)
    else:
        raw = str(artifact_or_path)
        if raw.lstrip().startswith(("{", "[")):
            pass
        else:
            path = Path(raw)
            if path.exists():
                raw = path.read_text(encoding="utf-8")
        doc = json.loads(raw)
    if not isinstance(doc, dict):
        raise ValueError("artifact must decode to a JSON object")
    if doc.get("schema_version") != expected_schema_version:
        raise ValueError(
            f"artifact.schema_version must be {expected_schema_version}, got {doc.get('schema_version')!r}"
        )
    return doc


def _coerce_inputs(
    assessment_table: Mapping[str, Any] | str | Path,
    map_doc: Mapping[str, Any] | str | Path,
    mar: Mapping[str, Any] | str | Path,
    profile_diff: Mapping[str, Any] | str | Path,
    bundle: Mapping[str, Any] | str | Path,
) -> tuple[dict[str, Any], dict[str, Any], dict[str, Any], dict[str, Any], dict[str, Any]]:
    return (
        _coerce_json_artifact(
            assessment_table, expected_schema_version="m15_assessment_table_v1"
        ),
        _coerce_json_artifact(map_doc, expected_schema_version="m15_map_v1"),
        _coerce_json_artifact(mar, expected_schema_version="m15_mar_v1"),
        _coerce_json_artifact(
            profile_diff, expected_schema_version="m15_profile_diff_report_v1"
        ),
        _coerce_json_artifact(bundle, expected_schema_version="m15_bundle_manifest_v1"),
    )


def _fmt_bool(value: Any) -> str:
    return "yes" if bool(value) else "no"


def _profile_label(doc: Mapping[str, Any]) -> str:
    requirements = doc.get("profile_requirements")
    if isinstance(requirements, Mapping):
        label = requirements.get("profile_label")
        if label:
            return str(label)
    return str(doc.get("jurisdiction_profile", "unknown"))


def _line(text: str) -> str:
    return str(text).strip()


def _list_line(label: str, value: str) -> str:
    return f"{label}: {value}"


@dataclass(frozen=True)
class _Section:
    title: str
    paragraphs: list[str]


def _overview_section(
    assessment: Mapping[str, Any],
    map_doc: Mapping[str, Any],
    mar: Mapping[str, Any],
    bundle: Mapping[str, Any],
) -> _Section:
    product_context = assessment.get("product_context") or {}
    summary = assessment.get("summary") or {}
    signoff = mar.get("signoff") or map_doc.get("signoff") or {}
    paragraphs = [
        _list_line("Jurisdiction profile", f"{_profile_label(assessment)} ({assessment.get('jurisdiction_profile', 'unknown')})"),
        _list_line("Sponsor", str(product_context.get("sponsor", "unknown"))),
        _list_line("Program", str(product_context.get("program_name", "unknown"))),
        _list_line("Compound", str(product_context.get("compound_name", "unknown"))),
        _list_line("Indication", str(product_context.get("indication", "unknown"))),
        _list_line("Assessment review status", str(assessment.get("review_status", "unknown"))),
        _list_line("MAP status", str(map_doc.get("document_status", "unknown"))),
        _list_line("MAR status", str(mar.get("document_status", "unknown"))),
        _list_line("Bundle status", str(bundle.get("bundle_status", "unknown"))),
        _list_line("Highest model impact", str(summary.get("highest_model_impact", "unknown"))),
        _list_line("Highest model risk", str(summary.get("highest_model_risk", "unknown"))),
        _list_line("Primary author", str(signoff.get("primary_author", "unknown"))),
        _list_line("QA reviewer", str(signoff.get("qa_reviewer", "unknown"))),
        _list_line("Approver", str(signoff.get("approver", "unknown"))),
        _list_line("Signoff status", str(signoff.get("status", "unknown"))),
    ]
    return _Section("Overview", paragraphs)


def _assessment_section(assessment: Mapping[str, Any]) -> _Section:
    paragraphs = [
        _list_line("Review status", str(assessment.get("review_status", "unknown"))),
        _list_line(
            "Assessment framing",
            str((assessment.get("profile_requirements") or {}).get("framing_text", "n/a")),
        ),
    ]
    for entry in assessment.get("entries") or []:
        if not isinstance(entry, Mapping):
            continue
        paragraphs.extend(
            [
                _line(
                    f"{entry.get('qoi_id', 'QOI')}: {entry.get('question_of_interest', 'unknown question')}"
                ),
                _list_line("Context of use", str(entry.get("context_of_use", "unknown"))),
                _list_line("Model influence", str(entry.get("model_influence", "unknown"))),
                _list_line("Model impact", str(entry.get("model_impact", "unknown"))),
                _list_line("Model risk", str(entry.get("model_risk", "unknown"))),
                _list_line(
                    "Consequence of wrong decision",
                    str(entry.get("consequence_of_wrong_decision", "unknown")),
                ),
                _list_line(
                    "Recommended reporting level",
                    str(entry.get("recommended_reporting_level", "unknown")),
                ),
                _list_line("Justification", str(entry.get("justification", "n/a"))),
            ]
        )
        refs: list[str] = []
        for ref in entry.get("evidence_refs") or []:
            if not isinstance(ref, Mapping):
                continue
            refs.append(
                f"{ref.get('path', 'unknown')} ({ref.get('role', 'evidence')}, sha256={ref.get('sha256', 'unknown')})"
            )
        if refs:
            paragraphs.append(_list_line("Evidence refs", "; ".join(refs)))
        paragraphs.append("")
    return _Section("Assessment Table Summary", [p for p in paragraphs if p != ""])


def _map_section(map_doc: Mapping[str, Any]) -> _Section:
    signoff = map_doc.get("signoff") or {}
    paragraphs = [
        _list_line("Document status", str(map_doc.get("document_status", "unknown"))),
        _list_line("Context of use", str(map_doc.get("context_of_use", "unknown"))),
        _list_line(
            "Planning framing",
            str((map_doc.get("profile_requirements") or {}).get("framing_text", "n/a")),
        ),
        _list_line("Primary author", str(signoff.get("primary_author", "unknown"))),
        _list_line("QA reviewer", str(signoff.get("qa_reviewer", "unknown"))),
        _list_line("Approver", str(signoff.get("approver", "unknown"))),
        _list_line("Signoff status", str(signoff.get("status", "unknown"))),
    ]
    for question in map_doc.get("questions") or []:
        if not isinstance(question, Mapping):
            continue
        paragraphs.append(
            _line(
                f"Question {question.get('question_id', 'unknown')}: {question.get('question_of_interest', 'unknown')}"
            )
        )
    for dataset in map_doc.get("planned_datasets") or []:
        if not isinstance(dataset, Mapping):
            continue
        paragraphs.append(
            _line(
                f"Planned dataset {dataset.get('dataset_id', 'unknown')}: {dataset.get('description', 'n/a')}"
            )
        )
    for method in map_doc.get("methods") or []:
        if not isinstance(method, Mapping):
            continue
        paragraphs.append(
            _line(f"Method {method.get('method_id', 'unknown')}: {method.get('description', 'n/a')}")
        )
    for criterion in map_doc.get("technical_acceptance_criteria") or []:
        if not isinstance(criterion, Mapping):
            continue
        paragraphs.append(
            _line(
                f"Criterion {criterion.get('criterion_id', 'unknown')}: {criterion.get('description', 'n/a')} | rule={criterion.get('acceptance_rule', 'n/a')}"
            )
        )
    return _Section("Model Analysis Plan", paragraphs)


def _mar_section(mar: Mapping[str, Any]) -> _Section:
    signoff = mar.get("signoff") or {}
    paragraphs = [
        _list_line("Document status", str(mar.get("document_status", "unknown"))),
        _list_line("Context of use", str(mar.get("context_of_use", "unknown"))),
        _list_line(
            "Results framing",
            str((mar.get("profile_requirements") or {}).get("framing_text", "n/a")),
        ),
        _list_line("Primary author", str(signoff.get("primary_author", "unknown"))),
        _list_line("QA reviewer", str(signoff.get("qa_reviewer", "unknown"))),
        _list_line("Approver", str(signoff.get("approver", "unknown"))),
        _list_line("Signoff status", str(signoff.get("status", "unknown"))),
    ]
    for result in mar.get("criterion_results") or []:
        if not isinstance(result, Mapping):
            continue
        paragraphs.append(
            _line(
                f"Criterion {result.get('criterion_id', 'unknown')}: {result.get('status', 'unknown')} | observed={result.get('observed_value', 'n/a')} | notes={result.get('notes', 'n/a')}"
            )
        )
    for question in mar.get("questions") or []:
        if not isinstance(question, Mapping):
            continue
        paragraphs.append(
            _line(
                f"Conclusion {question.get('question_id', 'unknown')}: {question.get('conclusion_status', 'unknown')} | {question.get('conclusion', 'n/a')}"
            )
        )
    deviations = mar.get("deviations") or []
    paragraphs.append(_list_line("Open deviations", str(len(deviations))))
    for limitation in mar.get("limitations") or []:
        paragraphs.append(_line(f"Limitation: {limitation}"))
    return _Section("Model Analysis Report", paragraphs)


def _profile_diff_section(profile_diff: Mapping[str, Any]) -> _Section:
    summary = profile_diff.get("summary") or {}
    paragraphs = [
        _list_line("Selected profile", str(profile_diff.get("selected_profile", "unknown"))),
        _list_line(
            "Compared profiles",
            ", ".join(str(x) for x in (profile_diff.get("compared_profiles") or [])),
        ),
        _list_line("Documents compared", str(summary.get("documents_compared", "unknown"))),
        _list_line("Profiles compared", str(summary.get("profiles_compared", "unknown"))),
        _list_line(
            "Sections with profile-specific diff",
            str(summary.get("sections_with_profile_specific_diff", "unknown")),
        ),
    ]
    for document in profile_diff.get("documents") or []:
        if not isinstance(document, Mapping):
            continue
        paragraphs.append(
            _line(
                f"Document {document.get('document_kind', 'unknown')}: common sections = {', '.join(str(x) for x in (document.get('common_sections') or []))}"
            )
        )
        for view in document.get("profile_views") or []:
            if not isinstance(view, Mapping):
                continue
            paragraphs.append(
                _line(
                    f"{view.get('profile_label', 'unknown')}: profile-only sections = {', '.join(str(x) for x in (view.get('profile_only_sections') or [])) or 'none'}"
                )
            )
    return _Section("Profile Diff Summary", paragraphs)


def _bundle_section(bundle: Mapping[str, Any]) -> _Section:
    integrity = bundle.get("integrity") or {}
    paragraphs = [
        _list_line("Bundle status", str(bundle.get("bundle_status", "unknown"))),
        _list_line(
            "Deterministic rerender verified",
            _fmt_bool(integrity.get("deterministic_re_render_verified")),
        ),
        _list_line("All hashes present", _fmt_bool(integrity.get("all_hashes_present"))),
        _list_line("Signoff roles complete", _fmt_bool(integrity.get("signoff_roles_complete"))),
        _list_line("Signoff roles distinct", _fmt_bool(integrity.get("signoff_roles_distinct"))),
    ]
    for missing in integrity.get("missing_required_roles") or []:
        paragraphs.append(_line(f"Missing required role: {missing}"))
    for artifact_name, artifact in sorted((bundle.get("artifacts") or {}).items()):
        if not isinstance(artifact, Mapping):
            continue
        paragraphs.append(
            _line(
                f"{artifact_name}: {artifact.get('path', 'unknown')} | sha256={artifact.get('sha256', 'unknown')} | role={artifact.get('role', 'unknown')}"
            )
        )
    return _Section("Bundle Integrity", paragraphs)


def _build_sections(
    assessment: Mapping[str, Any],
    map_doc: Mapping[str, Any],
    mar: Mapping[str, Any],
    profile_diff: Mapping[str, Any],
    bundle: Mapping[str, Any],
) -> list[_Section]:
    return [
        _overview_section(assessment, map_doc, mar, bundle),
        _assessment_section(assessment),
        _map_section(map_doc),
        _mar_section(mar),
        _profile_diff_section(profile_diff),
        _bundle_section(bundle),
    ]


def render_m15_publishable_markdown(
    assessment_table: Mapping[str, Any] | str | Path,
    map_doc: Mapping[str, Any] | str | Path,
    mar: Mapping[str, Any] | str | Path,
    profile_diff: Mapping[str, Any] | str | Path,
    bundle: Mapping[str, Any] | str | Path,
) -> str:
    assessment, map_doc, mar, profile_diff, bundle = _coerce_inputs(
        assessment_table, map_doc, mar, profile_diff, bundle
    )
    sections = _build_sections(assessment, map_doc, mar, profile_diff, bundle)
    out = ["# ICH M15 Publishable Report", ""]
    for section in sections:
        out.append(f"## {section.title}")
        out.append("")
        for paragraph in section.paragraphs:
            out.append(f"- {paragraph}")
        out.append("")
    return "\n".join(out).rstrip() + "\n"


def _wrap_paragraph(paragraph: str, *, width: int = 96) -> list[str]:
    if not paragraph:
        return [""]
    return textwrap.wrap(paragraph, width=width) or [paragraph]


def write_m15_publishable_pdf(
    pdf_path: str | Path,
    assessment_table: Mapping[str, Any] | str | Path,
    map_doc: Mapping[str, Any] | str | Path,
    mar: Mapping[str, Any] | str | Path,
    profile_diff: Mapping[str, Any] | str | Path,
    bundle: Mapping[str, Any] | str | Path,
) -> None:
    _require_matplotlib()
    _apply_pub_style()

    assessment, map_doc, mar, profile_diff, bundle = _coerce_inputs(
        assessment_table, map_doc, mar, profile_diff, bundle
    )
    sections = _build_sections(assessment, map_doc, mar, profile_diff, bundle)

    from matplotlib.backends.backend_pdf import PdfPages
    import matplotlib.pyplot as plt

    output_path = Path(pdf_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    pages_data: list[list[str]] = []
    current: list[str] = []
    max_lines = 36
    for section in sections:
        block = [section.title, ""]
        for paragraph in section.paragraphs:
            block.extend(_wrap_paragraph(paragraph))
            block.append("")
        if current and len(current) + len(block) > max_lines:
            pages_data.append(current)
            current = []
        if len(block) > max_lines and not current:
            chunk: list[str] = []
            for line in block:
                chunk.append(line)
                if len(chunk) >= max_lines:
                    pages_data.append(chunk)
                    chunk = []
            current = chunk
        else:
            current.extend(block)
    if current:
        pages_data.append(current)

    metadata = {
        "Title": "ICH M15 Publishable Report",
        "Creator": "nextstat.m15_report",
        "Producer": "NextStat (matplotlib)",
        "CreationDate": _FIXED_DT,
        "ModDate": _FIXED_DT,
    }

    total_pages = max(1, len(pages_data) + 1)
    with PdfPages(output_path, metadata=metadata) as pdf:
        cover = plt.figure(figsize=(8.27, 11.69))
        ax = cover.add_subplot(111)
        ax.axis("off")
        ax.text(
            0.0,
            0.97,
            "ICH M15 Publishable Report",
            ha="left",
            va="top",
            fontsize=20,
            weight="bold",
            color=_COL["ink"],
            transform=ax.transAxes,
        )
        ax.text(
            0.0,
            0.91,
            "\n".join(
                [
                    f"Jurisdiction profile: {_profile_label(assessment)} ({assessment.get('jurisdiction_profile', 'unknown')})",
                    f"Assessment review status: {assessment.get('review_status', 'unknown')}",
                    f"MAR status: {mar.get('document_status', 'unknown')}",
                    f"Bundle status: {bundle.get('bundle_status', 'unknown')}",
                    f"Primary author: {(mar.get('signoff') or {}).get('primary_author', 'unknown')}",
                    f"QA reviewer: {(mar.get('signoff') or {}).get('qa_reviewer', 'unknown')}",
                    f"Approver: {(mar.get('signoff') or {}).get('approver', 'unknown')}",
                ]
            ),
            ha="left",
            va="top",
            fontsize=11,
            family="monospace",
            color=_COL["ink"],
            transform=ax.transAxes,
            linespacing=1.3,
        )
        ax.text(
            0.0,
            0.02,
            f"Page 1/{total_pages}",
            ha="left",
            va="bottom",
            fontsize=8,
            color=_COL["muted"],
            transform=ax.transAxes,
        )
        pdf.savefig(cover)
        plt.close(cover)

        for page_index, lines in enumerate(pages_data, start=2):
            fig = plt.figure(figsize=(8.27, 11.69))
            ax = fig.add_subplot(111)
            ax.axis("off")
            y = 0.97
            for line in lines:
                if not line:
                    y -= 0.018
                    continue
                weight = "bold" if not line.startswith((" ", "-")) and ":" not in line[:32] else "normal"
                fontsize = 12 if weight == "bold" else 9.5
                ax.text(
                    0.0,
                    y,
                    line,
                    ha="left",
                    va="top",
                    fontsize=fontsize,
                    family="monospace" if ":" in line[:48] else None,
                    weight=weight,
                    color=_COL["ink"],
                    transform=ax.transAxes,
                )
                y -= 0.024 if fontsize > 10 else 0.02
            ax.text(
                0.0,
                0.02,
                f"Page {page_index}/{total_pages}",
                ha="left",
                va="bottom",
                fontsize=8,
                color=_COL["muted"],
                transform=ax.transAxes,
            )
            pdf.savefig(fig)
            plt.close(fig)


def _normalize_docx(path: Path) -> None:
    tmp_path = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as src:
        items = sorted((name, src.read(name)) for name in src.namelist())
    with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for name, data in items:
            info = zipfile.ZipInfo(name, date_time=_DOCX_ZIP_TIMESTAMP)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = 0
            info.comment = b""
            info.extra = b""
            dst.writestr(info, data)
    tmp_path.replace(path)


def write_m15_publishable_docx(
    docx_path: str | Path,
    assessment_table: Mapping[str, Any] | str | Path,
    map_doc: Mapping[str, Any] | str | Path,
    mar: Mapping[str, Any] | str | Path,
    profile_diff: Mapping[str, Any] | str | Path,
    bundle: Mapping[str, Any] | str | Path,
) -> None:
    _require_python_docx()
    assessment, map_doc, mar, profile_diff, bundle = _coerce_inputs(
        assessment_table, map_doc, mar, profile_diff, bundle
    )
    sections = _build_sections(assessment, map_doc, mar, profile_diff, bundle)

    from docx import Document
    from docx.shared import Inches, Pt

    output_path = Path(docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Courier New"
    normal_style.font.size = Pt(9.5)

    title = document.add_heading("ICH M15 Publishable Report", level=0)
    title.runs[0].font.name = "Arial"
    title.runs[0].font.size = Pt(18)

    for paragraph in _overview_section(assessment, map_doc, mar, bundle).paragraphs:
        document.add_paragraph(paragraph)

    for section_doc in sections[1:]:
        document.add_heading(section_doc.title, level=1)
        for paragraph in section_doc.paragraphs:
            document.add_paragraph(paragraph)

    core = document.core_properties
    core.author = "nextstat.m15_report"
    core.last_modified_by = "nextstat.m15_report"
    core.title = "ICH M15 Publishable Report"
    core.subject = "Frozen M15 reporting bundle"
    core.comments = "Deterministic publishable export"
    core.created = _FIXED_DT
    core.modified = _FIXED_DT

    document.save(output_path)
    _normalize_docx(output_path)


def _write_text(path: Path | None, content: str) -> None:
    if path is None:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def main(argv: Sequence[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Render a publishable M15 report from frozen artifacts.")
    sub = parser.add_subparsers(dest="command", required=True)

    render = sub.add_parser("render", help="Render Markdown/PDF/DOCX from frozen M15 artifacts.")
    render.add_argument("--assessment-table", required=True)
    render.add_argument("--map", required=True, dest="map_doc")
    render.add_argument("--mar", required=True)
    render.add_argument("--profile-diff", required=True)
    render.add_argument("--bundle", required=True)
    render.add_argument("--markdown")
    render.add_argument("--pdf")
    render.add_argument("--docx")

    args = parser.parse_args(argv)

    if args.command != "render":  # pragma: no cover
        parser.error("unknown command")

    if not any([args.markdown, args.pdf, args.docx]):
        parser.error("at least one of --markdown, --pdf, or --docx is required")

    markdown = render_m15_publishable_markdown(
        args.assessment_table,
        args.map_doc,
        args.mar,
        args.profile_diff,
        args.bundle,
    )
    _write_text(Path(args.markdown) if args.markdown else None, markdown)

    if args.pdf:
        write_m15_publishable_pdf(
            Path(args.pdf),
            args.assessment_table,
            args.map_doc,
            args.mar,
            args.profile_diff,
            args.bundle,
        )
    if args.docx:
        write_m15_publishable_docx(
            Path(args.docx),
            args.assessment_table,
            args.map_doc,
            args.mar,
            args.profile_diff,
            args.bundle,
        )
    return 0


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())
